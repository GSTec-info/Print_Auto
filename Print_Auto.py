import win32print # Importing the print module
import win32api # importing library to use print command
import os # Importing os module to use operating system commands
from datetime import date, datetime # Importing the datetime module and date, datetime methods
from docx import Document # Importing docx module and Document method
from docx.shared import Inches # Importing shared module and Inches method
from docx.enum.text import WD_ALIGN_PARAGRAPH # Importing the enum.text module and WD_ALIGN_PARAGRAPH method

############################################ GLOBAL VARIABLES ############################################## #

# Current date and time
dt_e_hr = datetime.now()
dt_e_hr_conv = date.strftime(dt_e_hr, '%d/%m/%Y at %H:%M:%S') # Convert to Brazilian format
dt_e_hr_conv_cmd = date.strftime(dt_e_hr, '%d/%m/%Y as %H:%M:%S') # Convert to Brazilian format (No special characters and accents)

# Path of 'Logo' file
pathLogo = r'path\to\logo.png'

# Path of color file
pathImageColor = r'path\to\image\color.png'

# Path to save the print file
pathSaveDocument = r'path\to\document.docx'

# Path to log path
pathLogCMD = 'path\\to\\Logs.log'

# Path of the print folder and the file to be printed
pathArchive = r"path\to\print\archive"
nameArchive = r"nameDocument.docx"

# Get list of printers
printer_list = win32print.EnumPrinters(2)

# Name of printers to check if it is active
defaultImp = 'default printer name'
findImp = 'name of the printer that will receive a printout'


############################################# CREATION OF DOCUMENT ############################################
# Document instance
make_doc = Document()

# Add Image 'Logo Arca' and Center (Header)
paragraph1 = make_doc.add_paragraph()
paragraph1.add_run().add_picture(pathLogo, width=Inches(1), height=Inches(0.8))
paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
make_doc.add_heading(' ', 2)

# Document Title
make_doc.add_heading('Auto Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

# Document Body
make_doc.add_heading(f'Generated at: {dt_e_hr_conv}', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
make_doc.add_heading(' ', 2)
make_doc.add_heading(' Print Quality: ', 5)
paragraph2 = make_doc.add_paragraph()
paragraph2.add_run().add_picture(pathImageColor, width=Inches(5.5), height=Inches(0.6))
paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the Document
make_doc.save(pathSaveDocument)



################################################# PRINT ######################################################

# Start execution log
os.system(f"cmd /c echo. >>{pathLogCMD}")
os.system(f"cmd /c echo --- Run on: {dt_e_hr_conv_cmd} >>{pathLogCMD}")

# Set the Printer 'IMP-QUALIDADE-COLOR' as default
for printer in printer_list:
    if printer[2] == findImp:
        win32print.SetDefaultPrinter(printer[2])
        # Log print log
        os.system(f"cmd /c echo --- Received printer: {printer[2]} >>{pathLogCMD}")
        os.system(f"cmd /c echo --- Printed Document: {nameArchive} >>{pathLogCMD}")
        print(printer[2])
        break

# Print the file
win32api.ShellExecute(0, "print", nameArchive, None, pathArchive, 0)
os.system('timeout 10')


# Set Printer 'IMP-TI' as default
for printer in printer_list:
    if printer[2] == defaultImp:
        win32print.SetDefaultPrinter(printer[2])
        print(printer[2])
        break

# End the script
os.system(f"cmd /c echo. >>{pathLogCMD}")
os.system(f"cmd /c echo --------------------End------------------- - >>{pathLogCMD}")

exit()