import win32print # Importing the print module
import win32api # importing library to use print command
import os # Importing os module to use operating system commands
from datetime import date, datetime # Importing the datetime module and date, datetime methods
from docx import Document # Importing docx module and Document method
from docx.shared import Inches # Importing shared module and Inches method
from docx.enum.text import WD_ALIGN_PARAGRAPH # Importing the enum.text module and WD_ALIGN_PARAGRAPH method

# Get current date and time
dt_e_hr = datetime.now()
dt_e_hr_conv = date.strftime(dt_e_hr, '%d/%m/%Y at %H:%M:%S') # Convert to Brazilian format
dt_e_hr_conv_cmd = date.strftime(dt_e_hr, '%d/%m/%Y as %H:%M:%S') # Convert to Brazilian format (No special characters and accents)



############################################# CREATION OF DOCUMENT ############################################
# Document instance
make_doc = Document()

# Add Image 'Logo' and Center (Header)
paragraph1 = make_doc.add_paragraph()
paragraph1.add_run().add_picture(r'path\to\image', width=Inches(1), height=Inches(0.8))
paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
make_doc.add_heading(' ', 2)

# Document Title
make_doc.add_heading('Auto Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

# Document Body
make_doc.add_heading(f'************ TEST ************', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
make_doc.add_heading(f'Generated at: {dt_e_hr_conv}', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
make_doc.add_heading(' ', 2)
make_doc.add_heading(' Print Quality: ', 5)
paragraph2 = make_doc.add_paragraph()
paragraph2.add_run().add_picture(r'path\to\image', width=Inches(5.5), height=Inches(0.6))
paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the Document
make_doc.save(r'path\to\save\archive\doc.docx')



############################################## # PRINT ################################################ ##########
# Start execution log
os.system(f"cmd /c echo. >>'path\\to\\log'")
os.system(f"cmd /c echo ************ TEST ************ >>path\\to\\log'")
os.system(f"cmd /c echo. >>path\\to\\log'")
os.system(f"cmd /c echo --- Run on: {dt_e_hr_conv_cmd} >>path\\to\\log'")

# Set the Printer 'IMP-QUALIDADE-COLOR' as default
printer_list = win32print.EnumPrinters(2)
printer = printer_list[0]
win32print.SetDefaultPrinter(printer[2])
print(printer[2])

os.system(f"cmd /c echo --- Received printer: {printer[2]} >>path\\to\\log'")

# Set the path of the print folder and the file to be printed
path = r"path\to\archive"
file = r"name\to\archive.docx"
os.system(f"cmd /c echo --- Printed Document: {file} >>path\\to\\log'")

# Print the file
win32api.ShellExecute(0, "print", file, None, path, 0)
os.system('timeout 10')


# Set Printer 'IMP-TI' as default
oldprinter = printer_list[4]
win32print.SetDefaultPrinter(OldPrinter[2])
print(OldPrinter[2])

# End the script
os.system(f"cmd /c echo. >>path\\to\\log'")
os.system(f"cmd /c echo --------------------End------------------- - >>path\\to\\log'")

exit()